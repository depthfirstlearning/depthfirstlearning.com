import re

from docx import Document


def get_caption_html(speaker, content, time, speaker_index):
    s = '<div class="caption">'
    s += '<div class="speakerTime">'
    s += '<div class="speaker speaker%d">%s</div>' % (speaker_index, speaker)
    s += '<div class="time">%s</div>' % time
    s += '</div>'
    s += '<div class="content">%s</div>' % content
    s += '</div>'
    return s


def convert(docx_path, output_path):
    m = re.compile('(.*): (\[\d{2}:\d{2}:\d{2}\]) (.*)')

    document = Document(docx_path)
    speakers = set()
    for paragraph in document.paragraphs:
        match = m.match(paragraph.text)
        if match is not None:
            speaker, _, _ = match.groups()
            speakers.add(speaker)
    speakers = sorted(speakers)

    captions = []
    speaker = None
    content = None
    time = None

    for paragraph in document.paragraphs:
        match = m.match(paragraph.text)
        if match is not None:
            if speaker is not None:
                speaker_index = speakers.index(speaker)
                captions.append(
                    get_caption_html(speaker, content, time, speaker_index)
                )
            speaker, time, content = match.groups()
            content = content.strip()
            time = time.strip()
            speaker = speaker.strip()
        elif paragraph.text and content is not None:
            content += paragraph.text.strip()

    speaker_index = speakers.index(speaker)
    captions.append(get_caption_html(speaker, content, time, speaker_index))
    with open(output_path, 'w') as f:
        f.write('<div class="captions">')
        for caption in captions:
            f.write(caption)
        f.write('</div>')
